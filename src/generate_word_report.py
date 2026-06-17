"""
Iki dilli (Ingilizce + Turkce) Word raporu uret.

Bu script proje boyunca olusan veri ve metrikleri kullanir:
- final_labeled_dataset.csv
- evaluation_metrics.json
- confusion_matrix.png

Uretilen dosyalar:
- docs/report_assets/*.png
- docs/AI_Network_Traffic_Classification_Report.docx
"""

from __future__ import annotations

import json
from pathlib import Path

import matplotlib
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


# Word raporu uretilirken ekrana pencere acilmamasi icin GUI istemeyen backend kullanilir.
matplotlib.use("Agg")

DATASET_PATH = Path("data/processed/final_labeled_dataset.csv")
EVALUATION_JSON_PATH = Path("models/evaluation_metrics.json")
CONFUSION_MATRIX_PATH = Path("models/confusion_matrix.png")
ASSETS_DIR = Path("docs/report_assets")
OUTPUT_DOCX_PATH = Path("docs/AI_Network_Traffic_Classification_Report.docx")

TRAFFIC_FILTERS = {
    "YouTube": 'ip.addr == 10.100.1.194 && (dns || quic || udp.port == 443)',
    "Browsing": 'ip.addr == 10.100.1.194 && (dns || tcp || tls)',
    "Download": 'ip.addr == 10.100.1.89 && (tcp || tls)',
}

ENGLISH_CODE_EXPLANATIONS = [
    (
        "Traffic-specific tshark filters",
        "These filters isolate only the packets that are most useful for each traffic type before CSV export.",
        [
            "YouTube: ip.addr == 10.100.1.194 && (dns || quic || udp.port == 443)",
            "Browsing: ip.addr == 10.100.1.194 && (dns || tcp || tls)",
            "Download: ip.addr == 10.100.1.89 && (tcp || tls)",
        ],
    ),
    (
        "PCAP to CSV export logic",
        "In traffic_signature_analysis.py, Python builds a tshark command and runs it with subprocess. This is how the project moves from packet capture files to structured CSV data.",
        [
            'command = [tshark_path, "-r", str(pcap_path), "-Y", display_filter, "-T", "fields"]',
            'for field in EXPORT_FIELDS:',
            '    command.extend(["-e", field])',
            'result = subprocess.run(command, capture_output=True, text=True)',
        ],
    ),
    (
        "CSV to final ML dataset logic",
        "In feature_extraction.py, exported CSV files are standardized into one beginner-friendly table. TCP and UDP ports are merged, protocol labels are simplified, and time_delta is calculated from packet timestamps.",
        [
            'standardized = pd.DataFrame({',
            '    "src_ip": ...',
            '    "dst_ip": ...',
            '    "src_port": coalesce_ports(df, "tcp.srcport", "udp.srcport"),',
            '    "dst_port": coalesce_ports(df, "tcp.dstport", "udp.dstport"),',
            '    "protocol": standardize_protocol(df),',
            '    "packet_length": ...',
            '    "time_delta": calculate_time_delta(df),',
            '    "label": ...',
            '})',
        ],
    ),
    (
        "Protocol ratio analysis logic",
        "In analyze_traffic_signatures.py, protocol ratios are measured by creating boolean masks. For example, ip.proto value 6 means TCP and 17 means UDP. Those values come from the packet data, while their meanings come from IP protocol standards.",
        [
            'if protocol_name == "tcp":',
            '    return (ip_proto == 6) | (tcp_src != "") | (tcp_dst != "")',
            'if protocol_name == "udp":',
            '    return (ip_proto == 17) | (udp_src != "") | (udp_dst != "")',
            'packet_ratio = packet_count / total_packets * 100',
        ],
    ),
    (
        "Model training logic",
        "In train_model.py, the final dataset is split into train and test subsets. Categorical features are encoded and a RandomForestClassifier is trained on the processed features.",
        [
            'x_train, x_test, y_train, y_test = train_test_split(...)',
            'preprocessor = ColumnTransformer([...])',
            'model = RandomForestClassifier(...)',
            'pipeline = Pipeline([("preprocessor", preprocessor), ("model", model)])',
            'pipeline.fit(x_train, y_train)',
        ],
    ),
    (
        "Model evaluation logic",
        "In evaluate_model.py, the saved model is loaded again, predictions are produced on the saved test dataset, and standard metrics are computed.",
        [
            'bundle = pickle.load(file_handle)',
            'y_pred = pipeline.predict(x_test)',
            'accuracy_score(y_true, y_pred)',
            'precision_score(y_true, y_pred, average="weighted")',
            'recall_score(y_true, y_pred, average="weighted")',
            'f1_score(y_true, y_pred, average="weighted")',
            'confusion_matrix(y_true, y_pred, labels=labels)',
        ],
    ),
]

TURKISH_CODE_EXPLANATIONS = [
    (
        "Trafik tipine ozel tshark filtreleri",
        "Bu filtreler, her trafik turu icin en anlamli paketleri secerek CSV export oncesi veriyi daraltir.",
        [
            "YouTube: ip.addr == 10.100.1.194 && (dns || quic || udp.port == 443)",
            "Browsing: ip.addr == 10.100.1.194 && (dns || tcp || tls)",
            "Download: ip.addr == 10.100.1.89 && (tcp || tls)",
        ],
    ),
    (
        "PCAP dosyasindan CSV olusturma mantigi",
        "traffic_signature_analysis.py icinde Python, tshark komutunu kurup subprocess ile calistirir. Boylece ham paket yakalama dosyalari yapilandirilmis CSV verisine donusur.",
        [
            'command = [tshark_path, "-r", str(pcap_path), "-Y", display_filter, "-T", "fields"]',
            'for field in EXPORT_FIELDS:',
            '    command.extend(["-e", field])',
            'result = subprocess.run(command, capture_output=True, text=True)',
        ],
    ),
    (
        "CSV dosyalarini son ML veri setine donusturme mantigi",
        "feature_extraction.py icinde export edilen CSV dosyalari daha sade bir tabloya cevrilir. TCP ve UDP portlari birlestirilir, protocol degerleri basitlestirilir ve time_delta zaman farki uretilir.",
        [
            'standardized = pd.DataFrame({',
            '    "src_ip": ...',
            '    "dst_ip": ...',
            '    "src_port": coalesce_ports(df, "tcp.srcport", "udp.srcport"),',
            '    "dst_port": coalesce_ports(df, "tcp.dstport", "udp.dstport"),',
            '    "protocol": standardize_protocol(df),',
            '    "packet_length": ...',
            '    "time_delta": calculate_time_delta(df),',
            '    "label": ...',
            '})',
        ],
    ),
    (
        "Protocol oranlarini hesaplama mantigi",
        "analyze_traffic_signatures.py icinde protocol oranlari boolean maskeler ile hesaplanir. Ornegin ip.proto degeri 6 ise TCP, 17 ise UDP anlamina gelir. Sayi paket verisinden gelir, anlami ise IP protocol standartlarindan gelir.",
        [
            'if protocol_name == "tcp":',
            '    return (ip_proto == 6) | (tcp_src != "") | (tcp_dst != "")',
            'if protocol_name == "udp":',
            '    return (ip_proto == 17) | (udp_src != "") | (udp_dst != "")',
            'packet_ratio = packet_count / total_packets * 100',
        ],
    ),
    (
        "Model egitme mantigi",
        "train_model.py icinde son veri seti train ve test olarak ayrilir. Kategorik feature'lar encode edilir ve RandomForestClassifier ile model egitilir.",
        [
            'x_train, x_test, y_train, y_test = train_test_split(...)',
            'preprocessor = ColumnTransformer([...])',
            'model = RandomForestClassifier(...)',
            'pipeline = Pipeline([("preprocessor", preprocessor), ("model", model)])',
            'pipeline.fit(x_train, y_train)',
        ],
    ),
    (
        "Model degerlendirme mantigi",
        "evaluate_model.py icinde kaydedilen model yeniden yuklenir, test verisi uzerinde tahmin uretilir ve temel metrikler hesaplanir.",
        [
            'bundle = pickle.load(file_handle)',
            'y_pred = pipeline.predict(x_test)',
            'accuracy_score(y_true, y_pred)',
            'precision_score(y_true, y_pred, average="weighted")',
            'recall_score(y_true, y_pred, average="weighted")',
            'f1_score(y_true, y_pred, average="weighted")',
            'confusion_matrix(y_true, y_pred, labels=labels)',
        ],
    ),
]


def load_dataset() -> pd.DataFrame:
    """Final labeled dataset dosyasini oku."""
    if not DATASET_PATH.exists():
        raise FileNotFoundError(f"Dataset file not found: {DATASET_PATH}")
    return pd.read_csv(DATASET_PATH, low_memory=False)


def load_evaluation_metrics() -> dict:
    """Model degerlendirme metriklerini JSON dosyasindan oku."""
    if not EVALUATION_JSON_PATH.exists():
        raise FileNotFoundError(f"Evaluation metrics file not found: {EVALUATION_JSON_PATH}")
    return json.loads(EVALUATION_JSON_PATH.read_text(encoding="utf-8"))


def create_label_distribution_figure(dataframe: pd.DataFrame) -> Path:
    """Notebook'taki label distribution mantigini PNG olarak kaydet."""
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    output_path = ASSETS_DIR / "label_distribution.png"

    counts = dataframe["label"].value_counts()
    plt.figure(figsize=(8, 5))
    counts.plot(kind="bar", color=["#4E79A7", "#F28E2B", "#59A14F"])
    plt.title("Label Distribution")
    plt.xlabel("Label")
    plt.ylabel("Record Count")
    plt.xticks(rotation=0)
    plt.tight_layout()
    plt.savefig(output_path, dpi=150)
    plt.close()
    return output_path


def create_protocol_distribution_figure(dataframe: pd.DataFrame) -> Path:
    """Notebook'taki protocol distribution mantigini PNG olarak kaydet."""
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    output_path = ASSETS_DIR / "protocol_distribution.png"

    counts = dataframe["protocol"].value_counts()
    plt.figure(figsize=(8, 5))
    counts.plot(kind="bar", color="#E15759")
    plt.title("Protocol Distribution")
    plt.xlabel("Protocol")
    plt.ylabel("Record Count")
    plt.xticks(rotation=0)
    plt.tight_layout()
    plt.savefig(output_path, dpi=150)
    plt.close()
    return output_path


def create_packet_length_figure(dataframe: pd.DataFrame) -> Path:
    """Notebook'taki average packet length by label grafigini PNG olarak kaydet."""
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    output_path = ASSETS_DIR / "packet_length_by_label.png"

    plt.figure(figsize=(8, 5))
    dataframe.groupby("label")["packet_length"].mean().plot(kind="bar", color="#76B7B2")
    plt.title("Average Packet Length by Label")
    plt.xlabel("Label")
    plt.ylabel("Average Packet Length")
    plt.xticks(rotation=0)
    plt.tight_layout()
    plt.savefig(output_path, dpi=150)
    plt.close()
    return output_path


def create_protocol_by_label_figure(dataframe: pd.DataFrame) -> Path:
    """Notebook'taki protocol by label stacked bar mantigini PNG olarak kaydet."""
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    output_path = ASSETS_DIR / "protocol_by_label.png"

    cross_tab = pd.crosstab(dataframe["label"], dataframe["protocol"])
    cross_tab.plot(kind="bar", stacked=True, figsize=(10, 6))
    plt.title("Protocol Distribution by Label")
    plt.xlabel("Label")
    plt.ylabel("Record Count")
    plt.xticks(rotation=0)
    plt.tight_layout()
    plt.savefig(output_path, dpi=150)
    plt.close()
    return output_path


def add_title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(18)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def add_paragraph(document: Document, text: str) -> None:
    document.add_paragraph(text)


def add_bullet_list(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_image(document: Document, image_path: Path, caption: str) -> None:
    if image_path.exists():
        document.add_picture(str(image_path), width=Inches(6.2))
        caption_paragraph = document.add_paragraph(caption)
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_code_lines(document: Document, lines: list[str]) -> None:
    """Word icinde basit bir kod blogu gorunumu olustur."""
    for line in lines:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)


def add_filter_list(document: Document) -> None:
    for traffic_name, filter_text in TRAFFIC_FILTERS.items():
        document.add_paragraph(f"{traffic_name}: {filter_text}", style="List Bullet")


def english_section(document: Document, dataframe: pd.DataFrame, metrics: dict, figures: dict[str, Path]) -> None:
    """Raporun Ingilizce bolumunu yaz."""
    add_heading(document, "English Report", level=1)

    add_heading(document, "1. Project Summary", level=2)
    add_paragraph(
        document,
        "This project was built to understand whether different application-level traffic patterns can be distinguished with packet-based features and classified with machine learning."
    )
    add_paragraph(
        document,
        "The workflow was designed as a step-by-step learning path: capture traffic with Wireshark, export packet fields with tshark, build a clean labeled dataset, analyze traffic signatures, train a model, and evaluate the final classifier."
    )

    add_heading(document, "2. What Was Collected", level=2)
    add_bullet_list(
        document,
        [
            "YouTube traffic from a controlled capture session",
            "Browsing traffic from normal web navigation",
            "Download traffic from a large file transfer session",
            "Each traffic type was labeled separately before dataset creation",
        ],
    )

    add_heading(document, "3. Traffic Filters Used in the Project", level=2)
    add_paragraph(
        document,
        "The following display filters were used to isolate traffic signatures for each capture. These filters were important because the project did not analyze the entire PCAP blindly; it focused on device-specific and protocol-specific behavior."
    )
    add_filter_list(document)

    add_heading(document, "4. What Was Analyzed in the Network Data", level=2)
    add_paragraph(
        document,
        "The project examined source and destination IPs, ports, protocol types, packet lengths, timing differences, DNS behavior, TLS behavior, and QUIC presence. These fields were first studied in Wireshark and later recreated in Python for repeatable analysis."
    )
    add_bullet_list(
        document,
        [
            "YouTube showed a strong QUIC + UDP signature",
            "Download traffic showed TCP + TLS dominance",
            "Browsing traffic showed mixed TCP, DNS, and encrypted web behavior",
        ],
    )

    add_heading(document, "5. How the Python Code Performed the Analysis", level=2)
    add_paragraph(
        document,
        "The project was not only analyzed manually in Wireshark. The same logic was implemented in Python step by step. The most important code stages are summarized below."
    )
    for title, explanation, lines in ENGLISH_CODE_EXPLANATIONS:
        add_heading(document, title, level=3)
        add_paragraph(document, explanation)
        add_code_lines(document, lines)

    add_heading(document, "6. Dataset Construction", level=2)
    add_paragraph(
        document,
        "Raw PCAP files were not used directly by the machine learning model. Instead, packet fields were exported into CSV format and converted into a cleaner final dataset with the following columns:"
    )
    add_bullet_list(
        document,
        [
            "src_ip",
            "dst_ip",
            "src_port",
            "dst_port",
            "protocol",
            "packet_length",
            "time_delta",
            "label",
        ],
    )
    add_paragraph(
        document,
        f"The final labeled dataset contains {len(dataframe):,} rows."
    )

    add_heading(document, "7. Exploratory Data Analysis", level=2)
    add_paragraph(
        document,
        "Exploratory data analysis was used to understand dataset quality and class behavior before training. Missing values, label imbalance, protocol frequency, and packet-length statistics were reviewed."
    )
    add_image(document, figures["label_distribution"], "Figure 1. Label distribution from the EDA notebook.")
    add_image(document, figures["protocol_distribution"], "Figure 2. Protocol distribution from the EDA notebook.")
    add_image(document, figures["packet_length"], "Figure 3. Average packet length by label.")
    add_image(document, figures["protocol_by_label"], "Figure 4. Protocol distribution by label.")

    add_heading(document, "8. Machine Learning Pipeline", level=2)
    add_paragraph(
        document,
        "A RandomForestClassifier was used as the baseline model. Categorical fields were encoded, numeric packet features were kept as numeric values, and the dataset was split into training and test subsets. The model was then evaluated on held-out test data."
    )

    add_heading(document, "9. Evaluation Metrics and Meaning", level=2)
    add_bullet_list(
        document,
        [
            "Accuracy: the overall percentage of correct predictions.",
            "Precision: how often the model is correct when it predicts a class.",
            "Recall: how many real examples of a class the model can find.",
            "F1-score: a balanced score between precision and recall.",
            "Confusion Matrix: shows which classes are confused with each other.",
        ],
    )
    add_paragraph(
        document,
        "Final evaluation metrics from the current saved run:"
    )
    add_bullet_list(
        document,
        [
            f"Accuracy: {metrics['accuracy']:.4f}",
            f"Weighted Precision: {metrics['precision_weighted']:.4f}",
            f"Weighted Recall: {metrics['recall_weighted']:.4f}",
            f"Weighted F1-score: {metrics['f1_weighted']:.4f}",
        ],
    )
    add_image(document, CONFUSION_MATRIX_PATH, "Figure 5. Confusion matrix from the saved evaluation output.")

    add_heading(document, "10. Step-by-Step Learning Outcomes", level=2)
    add_bullet_list(
        document,
        [
            "Learned what a PCAP file contains and how packet captures are inspected.",
            "Learned how to use Wireshark filters and tshark field exports.",
            "Learned the meaning of source IP, destination IP, ports, and protocol numbers.",
            "Learned how traffic signatures differ between streaming, browsing, and download sessions.",
            "Learned how to build a labeled CSV dataset for machine learning.",
            "Learned how to clean and explore the dataset with pandas and matplotlib.",
            "Learned how to train and evaluate a baseline classifier with scikit-learn.",
            "Learned how to interpret accuracy, precision, recall, F1-score, and confusion matrix results.",
        ],
    )

    add_heading(document, "11. Conclusion", level=2)
    add_paragraph(
        document,
        "This project demonstrates a full student-level workflow that connects networking and machine learning: packet capture, protocol analysis, feature extraction, dataset building, exploratory analysis, model training, and evaluation."
    )


def turkish_section(document: Document, dataframe: pd.DataFrame, metrics: dict, figures: dict[str, Path]) -> None:
    """Raporun Turkce bolumunu yaz."""
    document.add_section(WD_SECTION.NEW_PAGE)
    add_heading(document, "Turkce Rapor", level=1)

    add_heading(document, "1. Proje Ozeti", level=2)
    add_paragraph(
        document,
        "Bu proje, farkli uygulama trafiklerinin paket seviyesindeki ozellikler ile ayirt edilip edilemeyecegini ve makine ogrenmesi ile siniflandirilip siniflandirilamayacagini anlamak icin gelistirildi."
    )
    add_paragraph(
        document,
        "Calisma adim adim ogrenme mantigiyla ilerledi: Wireshark ile trafik yakalama, tshark ile alan export etme, temiz etiketli veri seti olusturma, trafik imzalarini analiz etme, modeli egitme ve son olarak modeli degerlendirme."
    )

    add_heading(document, "2. Toplanan Trafik Tipleri", level=2)
    add_bullet_list(
        document,
        [
            "Kontrollu bir oturumdan YouTube trafigi",
            "Normal web gezintisinden browsing trafigi",
            "Buyuk dosya transferinden download trafigi",
            "Her trafik tipi veri seti olusturulmadan once ayri etiketlendi",
        ],
    )

    add_heading(document, "3. Projede Kullanilan Trafik Filtreleri", level=2)
    add_paragraph(
        document,
        "Asagidaki display filter'lar, her trafik turunun ag imzasini ayiklamak icin kullanildi. Bu filtreler onemlidir cunku proje tum PCAP icerigini gelisiguzel degil, cihaza ve protokole ozel sekilde analiz etti."
    )
    add_filter_list(document)

    add_heading(document, "4. Ag Verisinde Neler Analiz Edildi", level=2)
    add_paragraph(
        document,
        "Projede source IP, destination IP, portlar, protocol tipleri, packet length, zaman farklari, DNS davranisi, TLS davranisi ve QUIC varligi incelendi. Bu alanlar once Wireshark'ta gozlemlendi, sonra ayni mantik Python ile tekrar uretildi."
    )
    add_bullet_list(
        document,
        [
            "YouTube trafiginde QUIC + UDP baskin bulundu",
            "Download trafiginde TCP + TLS baskin bulundu",
            "Browsing trafiginde TCP, DNS ve sifreli web davranisi birlikte goruldu",
        ],
    )

    add_heading(document, "5. Python Kodu Bu Analizi Nasil Yapti", level=2)
    add_paragraph(
        document,
        "Proje sadece Wireshark uzerinden manuel analiz ile ilerlemedi. Ayni mantik Python kodu icinde de adim adim kuruldu. En onemli kod asamalari asagida ozetlenmistir."
    )
    for title, explanation, lines in TURKISH_CODE_EXPLANATIONS:
        add_heading(document, title, level=3)
        add_paragraph(document, explanation)
        add_code_lines(document, lines)

    add_heading(document, "6. Veri Seti Olusturma", level=2)
    add_paragraph(
        document,
        "Ham PCAP dosyalari dogrudan modele verilmedi. Bunun yerine paket alanlari once CSV formatina donusturuldu, sonra daha temiz ve model icin uygun tek bir veri setine cevrildi."
    )
    add_bullet_list(
        document,
        [
            "src_ip",
            "dst_ip",
            "src_port",
            "dst_port",
            "protocol",
            "packet_length",
            "time_delta",
            "label",
        ],
    )
    add_paragraph(
        document,
        f"Olusturulan son etiketli veri seti toplam {len(dataframe):,} satir icermektedir."
    )

    add_heading(document, "7. Kesifsel Veri Analizi", level=2)
    add_paragraph(
        document,
        "Model egitiminden once veri kalitesini ve sinif davranislarini anlamak icin EDA yapildi. Eksik degerler, sinif dengesizligi, protocol dagilimi ve packet length istatistikleri incelendi."
    )
    add_image(document, figures["label_distribution"], "Sekil 1. EDA notebook'undan label dagilimi.")
    add_image(document, figures["protocol_distribution"], "Sekil 2. EDA notebook'undan protocol dagilimi.")
    add_image(document, figures["packet_length"], "Sekil 3. Label bazinda ortalama paket boyutu.")
    add_image(document, figures["protocol_by_label"], "Sekil 4. Label bazinda protocol dagilimi.")

    add_heading(document, "8. Makine Ogrenmesi Asamasi", level=2)
    add_paragraph(
        document,
        "Baslangic modeli olarak RandomForestClassifier kullanildi. Kategorik sutunlar sayisallastirildi, sayisal alanlar korunarak train/test ayrimi yapildi ve model ayrilan test verisi uzerinde sinandi."
    )

    add_heading(document, "9. Degerlendirme Metrikleri ve Anlamlari", level=2)
    add_bullet_list(
        document,
        [
            "Accuracy: Genel olarak tahminlerin ne kadarinin dogru oldugunu gosterir.",
            "Precision: Model bir sinifi tahmin ettiginde ne kadar dogru oldugunu gosterir.",
            "Recall: Gercek sinif orneklerinin ne kadarinin yakalandigini gosterir.",
            "F1-score: Precision ve recall arasinda dengeli bir olcudur.",
            "Confusion Matrix: Hangi sinifin hangi sinifla karistigini gosterir.",
        ],
    )
    add_paragraph(document, "Kaydedilen son calismadan alinan degerler:")
    add_bullet_list(
        document,
        [
            f"Accuracy: {metrics['accuracy']:.4f}",
            f"Agirlikli Precision: {metrics['precision_weighted']:.4f}",
            f"Agirlikli Recall: {metrics['recall_weighted']:.4f}",
            f"Agirlikli F1-score: {metrics['f1_weighted']:.4f}",
        ],
    )
    add_image(document, CONFUSION_MATRIX_PATH, "Sekil 5. Kaydedilen model degerlendirmesinden confusion matrix.")

    add_heading(document, "10. Adim Adim Neler Ogrenildi", level=2)
    add_bullet_list(
        document,
        [
            "PCAP dosyasinin ne oldugu ve paket yakalama mantigi ogrenildi.",
            "Wireshark filtreleri ve tshark alan export mantigi ogrenildi.",
            "Source IP, destination IP, port ve protocol numaralarinin anlami pekistirildi.",
            "Streaming, browsing ve download trafiginin farkli imzalar olusturdugu goruldu.",
            "Makine ogrenmesi icin etiketli CSV veri seti olusturma mantigi ogrenildi.",
            "Pandas ve matplotlib ile veri analizi yapma pratiği kazanildi.",
            "Scikit-learn ile temel bir siniflandirma modeli egitildi.",
            "Accuracy, precision, recall, F1-score ve confusion matrix yorumlama ogrenildi.",
        ],
    )

    add_heading(document, "11. Sonuc", level=2)
    add_paragraph(
        document,
        "Bu proje, networking ile machine learning alanlarini birlestiren uctan uca bir ogrenci projesi olarak tamamlandi. Paket yakalama, protocol analizi, feature extraction, veri seti olusturma, EDA, model egitimi ve model degerlendirme adimlari birlikte uygulanmis oldu."
    )


def build_document(dataframe: pd.DataFrame, metrics: dict, figures: dict[str, Path]) -> Document:
    """Word dokumanini olustur."""
    document = Document()
    add_title(document, "AI-Based Network Traffic Classification")
    subtitle = document.add_paragraph("Bilingual Project Report / Iki Dilli Proje Raporu")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    english_section(document, dataframe, metrics, figures)
    turkish_section(document, dataframe, metrics, figures)
    return document


def main() -> None:
    dataframe = load_dataset()
    metrics = load_evaluation_metrics()

    figures = {
        "label_distribution": create_label_distribution_figure(dataframe),
        "protocol_distribution": create_protocol_distribution_figure(dataframe),
        "packet_length": create_packet_length_figure(dataframe),
        "protocol_by_label": create_protocol_by_label_figure(dataframe),
    }

    document = build_document(dataframe, metrics, figures)
    OUTPUT_DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_DOCX_PATH)

    print(f"Word report created: {OUTPUT_DOCX_PATH}")
    print(f"Assets directory: {ASSETS_DIR}")


if __name__ == "__main__":
    main()
